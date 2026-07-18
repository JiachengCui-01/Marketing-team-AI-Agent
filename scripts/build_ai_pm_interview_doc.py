from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "Marketing_Agent_AI产品经理面试知识点与追问解答.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(35, 42, 52)
MUTED = RGBColor(92, 100, 112)
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"


def set_run_font(run, size=None, bold=None, color=None, east_asia="Microsoft YaHei"):
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    rfonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_para(doc, text="", *, size=11, bold=False, color=INK, after=6, before=0, align=None, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
        size, color = 16, BLUE
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(7)
        size, color = 13, BLUE
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
        size, color = 12, DARK_BLUE
    for run in p.runs:
        set_run_font(run, size=size, bold=True, color=color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=10.8, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=10.8, color=INK)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.20
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.2, color=INK)
    add_para(doc, "", after=2)


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        set_cell_shading(hdr[i], LIGHT_GRAY)
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run_font(r, size=9.6, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(str(value))
            set_run_font(r, size=9.3, color=INK)
    set_table_width(table, widths)
    add_para(doc, "", after=3)
    return table


def add_qa(doc, q, a, followups=None):
    p = add_para(doc, f"Q：{q}", size=10.8, bold=True, color=DARK_BLUE, after=3, before=3)
    p.paragraph_format.keep_with_next = True
    add_para(doc, f"A：{a}", size=10.6, color=INK, after=5)
    if followups:
        for item in followups:
            add_bullet(doc, f"追问答法：{item}", level=0)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name in ["Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    header = section.header.paragraphs[0]
    header.text = "Marketing Agent AI 产品经理面试复习文档"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def cover(doc):
    add_para(doc, "AI 产品经理面试复习资料", size=12, bold=True, color=MUTED, after=8)
    p = add_para(
        doc,
        "Marketing Agent 项目知识点与追问解答",
        size=24,
        bold=True,
        color=RGBColor(0, 0, 0),
        after=4,
    )
    p.paragraph_format.space_before = Pt(18)
    add_para(
        doc,
        "按 AI 产品经理岗位要求整理：项目定位、用户价值、AI 工作流、技术实现、数据安全、评估指标、面试追问与参考回答",
        size=12.5,
        color=RGBColor(55, 55, 55),
        after=14,
    )
    rows = [
        ("项目一句话", "面向企业营销团队的多 Agent AI 工作台，把内容生成、营销数据分析、市场研究、行业新闻摘要和营销图片生成整合到可登录、可追踪、可沉淀的产品原型中。"),
        ("目标岗位表达", "体现 AI PM 对场景拆解、MVP 取舍、AI 能力边界、可解释体验、数据闭环和工程协作的理解。"),
        ("核心用户", "市场营销、增长、内容运营、投放优化、品牌负责人以及需要营销洞察的业务团队。"),
        ("核心技术栈", "Python/FastAPI/SQLite/Anthropic SDK/Gemini Image/Next.js/React/TypeScript/Tailwind CSS。"),
        ("生成日期", date.today().strftime("%Y-%m-%d")),
    ]
    add_matrix(doc, ["字段", "说明"], rows, [1.35, 5.15])
    add_callout(
        doc,
        "面试开场建议",
        "不要把它介绍成“一个聊天机器人”。更好的说法是：我做的是一个面向营销团队工作流的 AI Agent 产品原型，重点不是单次生成，而是把任务理解、专家分工、工具调用、过程可视化、结果沉淀和文件交付串成完整闭环。",
    )
    doc.add_page_break()


def add_project_overview(doc):
    add_heading(doc, "1. 项目总览：你要先讲清楚什么", 1)
    add_heading(doc, "1.1 项目背景与业务问题", 2)
    for item in [
        "营销团队的日常工作链路分散：写内容、查趋势、看投放数据、做行业摘要、生成素材通常要在多个工具之间切换。",
        "通用聊天式 AI 容易停留在单点能力：能写一段文案，但难以稳定完成“研究-分析-创作-交付-沉淀”的业务闭环。",
        "企业用户需要可信与可追踪：AI 为什么这样输出、调用了什么专家、用了哪些工具、生成了哪些文件，都需要能被观察和复盘。",
        "多用户环境必须做数据隔离：会话、上传文件、生成产物、图片历史和新闻配置都应归属于具体用户。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "1.2 产品目标", 2)
    for item in [
        "把营销团队高频但碎片化的工作整合到一个 AI 工作台：内容、数据、研究、新闻、图片。",
        "用多 Agent 分工提升任务命中率，让不同专家拥有不同提示词、工具和输出规范。",
        "用 SSE trace 和 artifact 预览降低 AI 黑箱感，让用户看到系统正在做什么，并能拿到可下载交付物。",
        "用账号体系、会话持久化、文件管理和 SQLite 数据模型，把 Demo 产品化成可体验的 SaaS 原型。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "1.3 面试中的 30 秒版本", 2)
    add_callout(
        doc,
        "建议话术",
        "这个项目是一个企业营销团队 AI Agent 工作台。用户可以在 Web 端输入营销需求并上传文件，系统由 Orchestrator 判断任务类型，分派给内容、数据分析、市场研究等专家 Agent；后端通过 Anthropic 的工具调用、web search、code execution 和 Files API 完成研究与分析，同时支持 PDF、图片、新闻摘要等交付物。产品层面我重点解决的是营销场景的工作流闭环、AI 过程透明、结果沉淀、多用户隔离和 MVP 可扩展性。",
    )
    add_heading(doc, "1.4 功能地图", 2)
    add_matrix(
        doc,
        ["模块", "用户价值", "核心实现"],
        [
            ("聊天工作台", "自然语言提交营销任务，查看 AI 输出和执行过程", "Next.js 三栏布局、FastAPI、SSE、session 持久化"),
            ("内容生成", "生成 LinkedIn、小红书、邮件、广告、博客、PDF 一页纸", "Content Agent + platform skills + generate_pdf 工具"),
            ("营销分析", "上传数据后计算 CTR、CVR、CPC、CPA、ROAS 并给建议", "Analytics Agent + Files API + code_execution 沙箱"),
            ("市场研究", "获得近期趋势、竞品动作和来源引用", "Research Agent + web_search + 引用规范"),
            ("每日新闻", "按行业/语言/时间生成最近 24 小时新闻摘要", "news_configs、后台调度、手动刷新、取消保留窗口"),
            ("营销图片", "按淘宝、小红书、Amazon、Instagram 风格生成或再编辑图片", "Gemini image、ImageSkill、模板、cutout、history"),
            ("账号与数据", "多用户使用、资源隔离、历史沉淀", "PBKDF2 密码哈希、token、SQLite 外键、user_id 校验"),
        ],
        [1.25, 2.45, 2.8],
    )


def add_ai_pm_knowledge(doc):
    add_heading(doc, "2. AI 产品经理应掌握的核心知识点", 1)
    add_heading(doc, "2.1 场景与用户洞察", 2)
    for item in [
        "用户画像：营销经理关注效率和可交付结果；内容运营关注平台语气和多版本产出；投放/增长关注 KPI 计算和优化建议；管理者关注趋势、竞品和复盘沉淀。",
        "高频任务：生成平台文案、分析活动数据、整理行业信息、生成营销物料、制作可分享 PDF、持续追踪新闻。",
        "痛点优先级：工具切换成本高、AI 输出不可控、数据分析门槛高、资料难沉淀、生成结果缺少可交付形态。",
        "MVP 切入点：先做单人/多账号可用的工作台闭环，而不是一开始做复杂团队权限、审批流和企业知识库。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "2.2 AI Agent 产品设计", 2)
    for item in [
        "Orchestrator 不是简单路由器，而是任务理解、拆解、专家调用和最终综合输出的控制层。",
        "专家 Agent 的价值是把不同任务的提示词、工具、输出格式和边界分开管理，降低单 Agent 提示词膨胀和职责混乱。",
        "工具调用让模型从“只会生成文本”变成“能检索、能计算、能产出文件、能生成图片”的工作流执行者。",
        "AI PM 需要定义每个 Agent 的输入契约、输出契约、失败兜底、评估指标和用户可见状态。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "2.3 多模态与文件处理", 2)
    for item in [
        "结构化数据文件不直接塞进 Prompt，而是上传到模型文件系统，由 code_execution 沙箱读取和计算，降低上下文成本和幻觉风险。",
        "PDF、Word、CSV 可抽取文本作为上下文；图片转为模型可识别的 image block；营销图片生成可以使用上传图或 cutout 作为参考。",
        "文件上传需要白名单、大小限制、文件名清洗、用户归属、预览/下载权限校验。",
        "artifact 是 AI 产品很重要的概念：用户真正要的往往不是一段聊天，而是可预览、可下载、可复用的交付物。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "2.4 可解释性与可信 AI", 2)
    for item in [
        "SSE trace 把 started、delegating、orchestrator_response、specialist_done、assistant_delta、artifact_created、result、error 等事件展示给前端。",
        "研究类回答要求引用 URL 和日期，区分事实和推断；新闻摘要要求限定最近 24 小时窗口。",
        "数据分析要求用真实计算而不是模型心算；缺少列时说明缺失并跳过对应指标。",
        "失败不应该让产品崩溃：研究失败返回 unavailable，图片生成失败返回 graceful 200，新闻失败不覆盖上一次有效摘要。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "2.5 评价指标与产品成功标准", 2)
    add_matrix(
        doc,
        ["维度", "指标", "项目中如何落地/可扩展"],
        [
            ("效率", "完成一次营销任务的时间、用户从输入到交付物的步骤数", "统一工作台、文件上传、SSE、预览下载减少跨工具切换"),
            ("质量", "文案可用率、分析建议采纳率、研究引用有效率", "平台技能、KPI 计算规范、来源引用要求"),
            ("可信", "用户是否理解 AI 执行过程、失败率、可恢复率", "trace 面板、fallback complete 接口、失败消息持久化"),
            ("留存", "会话回访率、历史文件下载率、新闻任务开启率", "sessions、artifacts、news_configs 可沉淀这些行为"),
            ("成本", "每任务 token、web search 次数、图片生成成本", "orchestrator_response 记录 usage，可扩展为成本面板"),
            ("安全", "越权访问拦截率、上传校验通过率、敏感信息泄露风险", "user_id 查询、token、文件白名单、路径不外露"),
        ],
        [1.0, 2.0, 3.5],
    )


def add_architecture(doc):
    add_heading(doc, "3. 系统架构与数据流知识点", 1)
    add_heading(doc, "3.1 总体架构", 2)
    add_matrix(
        doc,
        ["层级", "技术/模块", "职责", "AI PM 讲法"],
        [
            ("前端", "Next.js、React、TypeScript、Tailwind", "聊天、会话、上传、预览、新闻、生图、执行追踪", "把 AI 能力包装成可操作的营销工作台，而不是裸 API"),
            ("API", "FastAPI、sse-starlette", "鉴权、REST、SSE、上传下载、调度入口", "承接前端交互和 Agent 执行的产品边界"),
            ("Agent", "Orchestrator、Content、Analytics、Research", "任务分发、专家执行、结果综合", "多专家分工对应真实营销岗位分工"),
            ("工具", "web_search、code_execution、Files API、PDF、Gemini Image", "检索、计算、文件交付、图片生成", "工具决定 AI 产品能否进入真实业务动作"),
            ("数据", "SQLite、WAL、外键、ContextVar", "用户、会话、消息、文件、产物、新闻、图片历史", "让结果沉淀、权限隔离和后续增长分析成为可能"),
        ],
        [0.9, 1.65, 2.1, 1.85],
    )
    add_heading(doc, "3.2 一次聊天任务的数据流", 2)
    for step in [
        "用户登录后在 ChatPanel 输入需求，可附带 CSV、Excel、JSON、PDF、DOCX 或图片。",
        "前端创建/复用 session，通过 /api/sessions/{id}/stream 建立 SSE，并把 token 放入查询参数兼容 EventSource。",
        "后端校验用户与 session 归属，读取上传文件。如果是数据文件，只传路径提示；如果是文本文件，抽取文本；如果是图片，构造成 image block。",
        "Orchestrator 把用户消息加入 Conversation，调用 Claude，根据工具 schema 选择 delegate_to_content_agent、delegate_to_analytics_agent 或 delegate_to_research_agent。",
        "专家 Agent 调用服务端工具或客户端工具，例如 code_execution 分析数据、web_search 查来源、generate_pdf 生成文件。",
        "后端通过 on_event 把执行事件推给前端，同时把用户消息、助手结果和 artifact 元数据持久化到 SQLite。",
        "前端把 assistant_delta 拼成流式消息，收到 artifact_created 后自动在右侧预览面板展示文件。",
    ]:
        add_number(doc, step)
    add_heading(doc, "3.3 为什么用 SSE 而不是 WebSocket", 2)
    add_para(
        doc,
        "当前业务主要是服务端向浏览器单向推送 AI 执行过程和输出增量，不需要高频双向实时协同。SSE 浏览器原生支持、实现简单、可自动重连，适合 AI 生成的流式结果。未来如果做多人协作、实时评论、共同编辑或任务控制台，再考虑 WebSocket。",
    )
    add_heading(doc, "3.4 数据库设计的产品意义", 2)
    for item in [
        "users/auth_tokens 支持从匿名 Demo 到真实账号体系，是 SaaS 产品化的第一步。",
        "sessions/messages 让一次营销任务可以持续推进和复盘，避免 AI 结果用完即丢。",
        "uploads/artifacts 把输入素材和输出产物都资产化，支撑预览、下载、历史沉淀。",
        "news_configs/news_summaries 体现主动服务能力，产品从被动问答扩展到定时信息服务。",
        "image_history/image_templates 体现 AIGC 资产管理，让用户能重编辑、查历史、复用模板。",
    ]:
        add_bullet(doc, item)


def add_module_details(doc):
    add_heading(doc, "4. 各模块知识点、追问与参考回答", 1)
    add_heading(doc, "4.1 Orchestrator 与多 Agent 编排", 2)
    add_para(
        doc,
        "Orchestrator 的系统提示词规定：不能自己写营销文案、不能自己算指标、不能自己声称外部事实，而必须调用对应专家。它通过 Anthropic Messages API 的 tool_use 机制在最多 12 轮内完成调度，并把最终结果综合成 Markdown。",
    )
    add_qa(
        doc,
        "为什么要做多 Agent，而不是一个大 Prompt 搞定？",
        "因为营销任务的能力边界很不同：文案要平台语气和创意规范，数据分析要真实计算，市场研究要外部来源和日期引用。拆成专家 Agent 后，每个专家可以拥有独立提示词、工具和输出格式，便于迭代、评估和降级。AI PM 角度看，这相当于把复杂业务流程拆成可管理的能力模块。",
        [
            "如果面试官追问成本：多 Agent 会增加调用次数，所以当前用“明显单任务就只委派一个专家”的规则控制成本。",
            "如果追问复杂度：MVP 先做三个专家，后续再按高频场景增加图表、品牌语气、审批等专家。",
        ],
    )
    add_qa(
        doc,
        "Orchestrator 如何避免幻觉？",
        "它的硬规则是：写文案必须委派内容专家，算数据必须委派分析专家，外部事实必须委派研究专家。研究专家使用 web_search 并要求引用来源，分析专家使用 code_execution 真实计算，内容专家使用平台技能限制格式。这不是完全消灭幻觉，但把高风险任务交给更合适的工具链处理。",
    )
    add_qa(
        doc,
        "如果专家失败怎么办？",
        "项目里有 unavailable_markdown 和 failed_specialists 机制。专家返回 unavailable/error 后，Orchestrator 不会反复重试同一专家，避免死循环；如果全部工具结果都是不可用，会直接把失败信息返回给用户。产品上要让失败可理解、可恢复，而不是沉默失败。",
    )

    add_heading(doc, "4.2 内容生成 Agent", 2)
    add_para(
        doc,
        "内容 Agent 是一个 B2B 营销文案专家。它根据 format、platform、task 选择 content_skills，如 LinkedIn、Twitter/X、小红书、Blog、Email、Ad Copy、PDF，并遵守每个平台的长度、结构、语气和禁忌。",
    )
    for item in [
        "LinkedIn：面向 B2B 决策者，短段落、业务问题、结果导向、轻 CTA。",
        "小红书：经验分享、口语化、移动端短段落、自然使用 emoji、种草感。",
        "Email：Subject、Preheader、单一 CTA、正文短。",
        "Ad Copy：2-3 个变体，每个包含 Headline、Description、CTA，并限制长度。",
        "PDF：调用 generate_pdf 工具，生成多章节可下载文件，而不是把全部 PDF 内容贴回聊天。",
    ]:
        add_bullet(doc, item)
    add_qa(
        doc,
        "内容生成如何体现产品经理的场景理解？",
        "不是让模型随便写，而是把不同渠道的内容规范沉淀成技能库。比如 LinkedIn 和小红书的受众、语言风格、段落结构完全不同。AI PM 要把平台经验产品化成可复用规则，才能稳定生成更贴近业务的内容。",
    )
    add_qa(
        doc,
        "为什么内容 Agent 可以在信息缺失时做合理假设？",
        "营销工作里用户经常只给一个粗略需求，如果每次都追问会破坏效率。当前策略是缺少 audience、product、tone 时做一个合理假设并说明。产品上这是“降低启动成本”和“保证可控性”的平衡。",
    )

    add_heading(doc, "4.3 数据分析 Agent", 2)
    add_para(
        doc,
        "Analytics Agent 使用 Anthropic Files API 把 CSV/Excel/JSON 上传到代码执行沙箱，再由 pandas/openpyxl 读取文件。系统提示词要求先检查列和类型，再计算 CTR、CVR、CPC、CPA、ROAS，以及按渠道/活动的趋势和行动建议。",
    )
    add_matrix(
        doc,
        ["KPI", "公式", "产品解释"],
        [
            ("CTR", "clicks / impressions", "衡量曝光转点击效率，反映素材、定向和标题吸引力"),
            ("CVR", "conversions / clicks", "衡量点击后转化效率，反映落地页、优惠、用户匹配度"),
            ("CPC", "spend / clicks", "每次点击成本，常用于投放成本控制"),
            ("CPA", "spend / conversions", "每次转化成本，评估获客或转化效率"),
            ("ROAS", "revenue / spend", "广告收入回报，衡量投放是否值得继续加预算"),
        ],
        [0.9, 1.45, 4.15],
    )
    add_qa(
        doc,
        "为什么不用模型直接读取 CSV 文本并回答？",
        "直接把原始数据塞进 Prompt 成本高、上下文容易爆、模型还可能心算出错。代码执行沙箱能用 pandas 做真实计算，原始数据不进入提示词，只把聚合结果和结论返回，更适合真实营销数据分析场景。",
    )
    add_qa(
        doc,
        "如果数据缺列怎么办？",
        "系统提示词明确要求：不要发明数据，缺少列就说明并跳过对应指标。比如没有 revenue 就不能算 ROAS，没有 conversions 就不能算 CPA/CVR。AI PM 要关注数据可用性边界，而不是强行给完整结论。",
    )
    add_qa(
        doc,
        "分析结果如何评估好坏？",
        "不仅看是否有均值或总计，更看是否输出 3-5 个可行动发现和 3 个具体建议。比如“LinkedIn ROAS 环比提升 18%，建议增加预算并复用素材结构”比“平均 ROAS 是 9.4”更有业务价值。",
    )

    add_heading(doc, "4.4 市场研究与新闻摘要", 2)
    add_para(
        doc,
        "Research Agent 使用 web_search，最多 3 次搜索，要求 3-5 个可靠来源、每个事实带 URL 和日期、区分事实与推断。News 模块复用 Research Agent，为用户配置的行业生成最近 24 小时摘要，并支持中英文、brief/detailed、时区和定时调度。",
    )
    add_qa(
        doc,
        "为什么新闻摘要要限定最近 24 小时？",
        "用户打开每日新闻功能时期待的是“今天有什么新变化”，如果混入旧材料会破坏信任。项目在 build_task 里明确 window_start/window_end，并要求只包含窗口内发布内容；如果没有足够新闻，就直接说明，而不是用旧新闻凑数。",
    )
    add_qa(
        doc,
        "新闻生成失败为什么不覆盖旧摘要？",
        "这是产品可靠性设计。用户宁愿看到上一份有效摘要，也不希望有效内容被一条失败消息覆盖。项目通过 _research_failed 检测不可用结果，失败时抛 NewsGenerationError，不写入 news_summaries，也不更新 last_run_at。",
    )
    add_qa(
        doc,
        "新闻取消为什么不是立刻删除所有数据？",
        "当前设计是软取消：停止调度，但保留上一份摘要直到下一天配置的 summary_time。这样用户不会因为误操作立刻丢失内容；过了 revert_at 再清空配置和摘要，回到未激活状态。这体现了可撤销和状态过渡设计。",
    )

    add_heading(doc, "4.5 营销图片生成", 2)
    add_para(
        doc,
        "图片模块使用 Gemini image 模型。它抽象出 ImageSkill，覆盖小红书、淘宝、Amazon、Instagram 和通用营销图，每个技能包含平台别名、默认比例、风格规则、背景指导和负面提示。用户可上传参考图、选择原图或 cutout、选择模板、生成、查看历史和再编辑。",
    )
    add_matrix(
        doc,
        ["风格", "默认比例", "核心约束"],
        [
            ("小红书", "3:4", "暖光、生活场景、种草氛围、顶部留标题空间"),
            ("淘宝/天猫", "1:1", "白底或浅色棚拍、商品居中、细节清晰"),
            ("Amazon", "1:1", "纯白背景、仅商品、合规主图"),
            ("Instagram", "4:5", "杂志感、统一色调、信息流友好"),
            ("通用", "1:1", "干净、专业、突出主体"),
        ],
        [1.15, 1.0, 4.35],
    )
    add_qa(
        doc,
        "图片生成为什么需要平台风格技能？",
        "因为不同平台的图片规范完全不同。Amazon 主图强调合规白底和商品占比，小红书强调生活化种草，Instagram 强调审美和信息流。把这些规则结构化后，用户不用自己写复杂 Prompt，也减少生成结果偏离平台的概率。",
    )
    add_qa(
        doc,
        "为什么要做 cutout 背景移除？",
        "电商和营销图片常需要把商品从原始环境中提取出来，再放入新的平台化场景。项目先用 Claude vision 判断上传图是 object 还是 screenshot，只有物体照才尝试 rembg cutout；如果不确定默认保留完整截图，这是非破坏性策略。",
    )
    add_qa(
        doc,
        "图片生成失败如何处理？",
        "generate_image 永不抛异常给路由层，而是返回 ok=false、unavailable=true 和可读消息。比如缺 GEMINI_API_KEY、配额不足、安全拦截、空响应都会优雅降级，保证其他功能不受影响。",
    )

    add_heading(doc, "4.6 前端工作台与用户体验", 2)
    add_para(
        doc,
        "前端主页面是三栏工作台：左侧会话/分组，中间聊天、新闻或图片工作区，右侧执行 trace 与文件预览。它不是营销落地页，而是面向高频使用者的操作台。",
    )
    add_qa(
        doc,
        "为什么做三栏布局？",
        "营销任务往往需要同时看历史、输入需求、观察执行过程和预览交付物。三栏布局把这些高频对象并列呈现，减少来回跳转。左侧负责任务管理，中间负责工作流，右侧负责可信度和产物。",
    )
    add_qa(
        doc,
        "流式失败为什么要 fallback 到 complete 接口？",
        "SSE 可能因为网络、代理或浏览器连接中断失败。前端在 stream error 后先尝试从服务端历史恢复最近 assistant 消息，如果没有，再调用非流式 completeSession 兜底。这提升了任务完成率，而不是把错误直接抛给用户。",
    )
    add_qa(
        doc,
        "PreviewPanel 的产品意义是什么？",
        "它把执行 trace 和 artifact 预览放在聊天之外，让主对话保持清晰，同时让用户看到 AI 的过程和文件结果。AI PM 角度，这是把“可解释性”和“交付物管理”产品化。",
    )

    add_heading(doc, "4.7 鉴权、安全与数据隔离", 2)
    add_para(
        doc,
        "项目有完整账号注册登录：账号可以是邮箱或中国大陆手机号，密码用 PBKDF2-SHA256 加盐 200000 轮哈希，token 有 14 天 TTL。所有业务资源都通过 user_id 校验归属。",
    )
    add_qa(
        doc,
        "多用户隔离是怎么做的？",
        "数据库表中 sessions、groups、uploads、artifacts、news_configs、image_history 等都包含 user_id，路由读取资源时必须同时匹配资源 id 和当前 user_id。测试里覆盖了跨用户访问会话、图片历史和文件的 404 拦截。",
    )
    add_qa(
        doc,
        "上传文件有哪些安全策略？",
        "只允许白名单扩展名，包括 CSV、Excel、JSON、PDF、DOCX 和常见图片；限制 10MB；清洗文件名；保存时使用 UUID 前缀；下载和预览都需要鉴权并检查 user_id。接口不会把真实 path 暴露给前端。",
    )
    add_qa(
        doc,
        "当前安全方案有什么不足？",
        "MVP 还没有企业级 RBAC、团队空间、审计日志、病毒扫描、内容安全扫描、对象存储签名 URL 和细粒度权限。面试中要主动承认这是 MVP 取舍，并说明未来会按企业化路线补齐。",
    )


def add_interview_bank(doc):
    add_heading(doc, "5. 高频面试追问题库与回答", 1)
    sections = [
        (
            "5.1 产品定位与用户价值",
            [
                ("这个项目解决的核心问题是什么？", "核心问题是营销团队工作流碎片化和通用 AI 输出不可沉淀。它把内容、数据、研究、新闻和图片放进一个可登录、可追踪、可交付的 AI 工作台，让用户从“问 AI 一句话”升级为“完成一次营销任务”。"),
                ("你的目标用户是谁？", "第一批用户是中小团队或企业内部营销团队，包括内容运营、投放优化、增长经理、品牌负责人。选择他们是因为营销工作高频、内容和数据都多、对生成式 AI 接受度高，且有明确可交付成果。"),
                ("为什么不先做通用办公 AI？", "通用办公范围太大，评估标准模糊。营销场景有清晰任务链、平台规范、常见 KPI 和文件交付需求，适合做垂直场景 MVP，也更容易证明价值。"),
                ("项目的北极星指标是什么？", "可设为“每周成功完成的营销任务数”或“AI 交付物被预览/下载/继续编辑的次数”。它比登录次数更贴近真实价值，因为用户最终要的是可用结果。"),
            ],
        ),
        (
            "5.2 MVP 与优先级",
            [
                ("你如何确定 MVP 范围？", "我优先选择能形成闭环的能力：账号、会话、上传、Agent 编排、核心专家、SSE 过程、artifact 交付。没有先做团队协作、计费、复杂权限和知识库，因为那些需要在单用户闭环验证后再扩展。"),
                ("如果只保留一个功能，你保留什么？", "保留“聊天工作台 + 多 Agent 编排 + artifact/trace”。因为这是产品差异化的核心，其余新闻、图片等都可以作为后续场景模块扩展。"),
                ("如何排序后续需求？", "按用户价值、使用频次、实现成本和风险排序。短期优先稳定性、成本展示、重试取消、模板；中期做品牌语气库、图表报告、会话搜索；长期做团队空间、权限、审批、数据连接器。"),
            ],
        ),
        (
            "5.3 AI 能力与模型边界",
            [
                ("你如何降低大模型幻觉？", "根据任务风险选择工具：外部事实必须 web_search 并引用，数据分析必须 code_execution，文件内容先抽取或上传沙箱，缺列不编。再通过 trace 和失败兜底让用户知道系统做了什么。"),
                ("为什么选 Claude/Anthropic？", "项目使用 Anthropic SDK 是因为它提供 Messages API、tool_use、web_search、code_execution、Files API 等能力，适合做多 Agent 与工具工作流。PM 表达上不是“迷信某模型”，而是模型和工具能力匹配产品目标。"),
                ("如果模型成本太高怎么办？", "可以按任务分层：Orchestrator 用强模型，简单内容改用低成本模型；研究和分析按需开启；缓存新闻摘要和模板结果；记录 usage 做成本看板；对图片生成设置额度。"),
                ("如何做 Agent 评估？", "离线评估看路由准确率、工具调用正确率、输出格式合规率、引用有效率、KPI 计算准确率；在线评估看用户重试率、采纳率、下载率、会话完成率和负反馈。"),
            ],
        ),
        (
            "5.4 数据与分析",
            [
                ("营销分析为什么重点算 CTR、CVR、CPC、CPA、ROAS？", "它们覆盖曝光到点击、点击到转化、成本效率和收入回报，是投放和增长团队最常见的基础指标，能快速给出优化方向。"),
                ("如何保证分析结果可信？", "让代码执行工具读取原始文件并用 pandas 计算；先检查列和类型；缺失字段明确说明；只报告聚合和计算结果；测试覆盖分析 Agent 的错误处理。"),
                ("如果数据质量很差怎么办？", "产品应先做数据质量提示：缺失值、异常值、重复行、字段映射、时间粒度不一致。当前 MVP 主要靠 Agent 检查列和 dtype，后续可做上传后的自动数据诊断面板。"),
            ],
        ),
        (
            "5.5 安全、隐私与合规",
            [
                ("项目是否能直接给企业使用？", "当前更像产品化 MVP，不是完整企业版。已有账号、token、密码哈希、用户隔离、文件白名单，但企业化还需要团队权限、审计、加密存储、对象存储、内容安全、合规模块和 SLA。"),
                ("为什么 SSE token 放 query 参数有风险吗？", "EventSource 原生设置自定义 Header 不方便，所以项目用 query token 兼容 SSE。风险是 URL 日志可能记录 token，生产环境应缩短 token、使用 httpOnly cookie 或服务端 session，并控制日志脱敏。"),
                ("SQLite 能支撑生产吗？", "SQLite 适合单进程 MVP，部署简单、成本低；项目也明确 scheduler 假设单 worker。规模上来后应迁移 PostgreSQL，并为新闻调度和任务执行加分布式锁/队列。"),
            ],
        ),
        (
            "5.6 前端体验与商业化",
            [
                ("为什么不是普通聊天界面？", "营销工作需要多任务管理、过程观察和文件交付。普通聊天界面会把 trace、预览、历史都挤在一条消息流里，三栏工作台更适合高频业务用户。"),
                ("如何商业化？", "可以按 seat、任务量、模型用量、图片生成额度、高级研究、团队知识库、数据连接器和模板库收费。营销团队愿意为节省内容和分析时间、提升产出质量买单。"),
                ("如何让用户形成习惯？", "通过每日新闻、历史会话、模板、品牌语气库、图片历史和 artifact 管理，把产品从一次性生成工具变成每天打开的营销工作台。"),
            ],
        ),
    ]
    for title, qas in sections:
        add_heading(doc, title, 2)
        for q, a in qas:
            add_qa(doc, q, a)


def add_limitations_roadmap(doc):
    add_heading(doc, "6. 当前不足、风险与路线图", 1)
    add_heading(doc, "6.1 当前限制", 2)
    for item in [
        "单进程假设：新闻调度器在 FastAPI lifespan 中启动，多 worker 会重复执行，需要分布式锁或任务队列。",
        "数据库扩展性：SQLite 适合 MVP，但高并发、多租户和复杂查询需要 PostgreSQL。",
        "权限体系：目前是个人账号级隔离，没有团队空间、角色权限、共享会话和审批流。",
        "评估体系：代码中已有测试，但还没有系统化 Prompt/Agent 离线评测集和人工反馈闭环。",
        "成本治理：已能记录部分 token usage，但还没有用户可见的成本预算、配额和告警。",
        "内容安全：图片与文案生成的安全、版权、品牌合规目前主要依赖模型侧和提示词，产品侧策略还不完整。",
        "数据连接：目前以文件上传为主，未接入广告平台、CRM、GA、BI 等真实业务系统。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "6.2 路线图", 2)
    add_matrix(
        doc,
        ["阶段", "重点", "为什么重要"],
        [
            ("短期", "任务取消/重试、错误提示优化、成本展示、模板推荐、会话搜索", "提升可控感和日常可用性"),
            ("中期", "品牌语气库、Prompt 模板、图表报告、数据质量检测、更多文件产物", "从通用 Agent 变成营销团队专属工具"),
            ("长期", "团队空间、RBAC、审批流、审计日志、广告/CRM/BI 连接器", "进入企业协作和商业化场景"),
            ("商业化", "按席位/用量/图片额度/高级研究/连接器收费", "形成清晰收入模型并控制推理成本"),
        ],
        [1.0, 3.0, 2.5],
    )
    add_heading(doc, "6.3 如果让你重新做，你会改什么", 2)
    for item in [
        "先定义更明确的事件模型和任务状态机，便于前端展示、重试、取消和审计。",
        "把 Agent 执行从 API 请求中拆到后台任务队列，解决长任务、失败重试和水平扩展问题。",
        "更早加入评测集：路由准确率、工具调用成功率、KPI 计算准确率、引用有效率、内容合规率。",
        "把文件和 artifact 存储迁到对象存储，数据库只存元数据和签名访问策略。",
        "为营销图片加入版权/品牌安全检查、尺寸裁剪、批量生成和 A/B 素材管理。",
    ]:
        add_bullet(doc, item)


def add_resume_and_story(doc):
    add_heading(doc, "7. 简历表达与 STAR 面试故事", 1)
    add_heading(doc, "7.1 简历项目描述", 2)
    for item in [
        "设计并落地企业营销 Multi-Agent AI 工作台，覆盖内容生成、营销数据分析、市场研究、行业新闻摘要和营销图片生成等场景。",
        "将 AI Agent Demo 产品化为 Web SaaS 原型，补齐注册登录、多用户隔离、会话分组、文件上传、SSE 执行追踪、artifact 预览下载和历史沉淀。",
        "设计 Orchestrator + 专家 Agent 架构，结合 web_search、code_execution、Files API、PDF 生成和图像生成工具，提升营销任务完成度和可解释性。",
        "围绕可信 AI 设计来源引用、真实计算、失败兜底、trace 可视化和结果持久化机制，降低黑箱感和业务使用风险。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "7.2 STAR 故事", 2)
    add_matrix(
        doc,
        ["STAR", "内容"],
        [
            ("Situation", "营销团队在内容、数据、研究和素材生成之间频繁切换工具，通用 AI 难以沉淀结果和解释过程。"),
            ("Task", "做一个能展示 AI PM 能力的产品化 MVP，不只是聊天，而是可登录、可追踪、可交付的营销 AI 工作台。"),
            ("Action", "拆出 Orchestrator、内容、分析、研究等 Agent；设计文件上传、SSE、artifact、新闻调度、图片生成、用户隔离和测试覆盖。"),
            ("Result", "形成完整全栈原型，覆盖从需求输入、专家分工、工具执行、过程展示到文件交付和历史沉淀的闭环，可用于演示 AI 产品设计与工程落地能力。"),
        ],
        [1.0, 5.5],
    )
    add_heading(doc, "7.3 面试收尾话术", 2)
    add_callout(
        doc,
        "建议话术",
        "这个项目让我理解到，AI 产品经理不能只会写 Prompt，更要把业务任务拆成可执行工作流，定义模型和工具的边界，设计失败兜底和评估指标，再把结果包装成用户能信任、能复用、能沉淀的产品体验。",
    )


def add_glossary(doc):
    add_heading(doc, "8. 术语速查", 1)
    add_matrix(
        doc,
        ["术语", "解释", "项目对应"],
        [
            ("Agent", "能根据目标调用模型和工具完成任务的执行单元", "Content/Analytics/Research Agent"),
            ("Orchestrator", "负责理解、拆解、分发和综合的控制层", "orchestrator.py"),
            ("Tool Calling", "模型输出工具调用请求，由系统执行后返回结果", "delegate_*、generate_pdf"),
            ("SSE", "Server-Sent Events，服务端向前端单向推送事件", "/stream、PreviewPanel trace"),
            ("Artifact", "AI 生成的可预览/下载产物", "PDF、image、cutout"),
            ("Code Execution", "模型可使用代码沙箱执行真实计算", "Analytics Agent"),
            ("Files API", "把文件上传给模型工具环境使用", "container_upload"),
            ("Prompt Skill", "结构化的场景规则和输出契约", "content_skills、image_skills"),
            ("Graceful Degradation", "依赖失败时不崩溃，返回可理解状态", "Image Unavailable、Research Unavailable"),
            ("Multi-Tenant Isolation", "多用户资源隔离", "所有资源按 user_id 查询"),
        ],
        [1.3, 2.35, 2.85],
    )


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    cover(doc)
    add_project_overview(doc)
    add_ai_pm_knowledge(doc)
    add_architecture(doc)
    add_module_details(doc)
    add_interview_bank(doc)
    add_limitations_roadmap(doc)
    add_resume_and_story(doc)
    add_glossary(doc)
    doc.save(OUT)
    print("DOCX_CREATED")


if __name__ == "__main__":
    build()
