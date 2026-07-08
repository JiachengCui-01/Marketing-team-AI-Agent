from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "Marketing_Agent_新版本项目技术文档.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 98, 110)
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9E2F3"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"


def set_run_font(run, size=None, bold=None, color=None, east_asia="Microsoft YaHei"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_para(doc, text="", style=None, size=11, bold=False, color=None, after=6, before=0, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    if level == 1:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
        size, color = 16, BLUE
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        size, color = 13, BLUE
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        size, color = 12, DARK_BLUE
    for run in p.runs:
        set_run_font(run, size=size, bold=True, color=color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run, size=10.8)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run, size=10.8)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.2, color=RGBColor(45, 52, 60))
    add_para(doc, "", after=2)


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, LIGHT_GRAY)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run_font(r, size=9.6, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.10
            r = p.add_run(value)
            set_run_font(r, size=9.4)
    set_table_width(table, widths)
    add_para(doc, "", after=3)
    return table


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    header = section.header.paragraphs[0]
    header.text = "Marketing Agent 新版本项目技术文档"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def cover(doc):
    add_para(doc, "项目技术文档", size=12, bold=True, color=MUTED, after=8)
    p = add_para(
        doc,
        "Marketing Agent 新版本项目总结",
        size=25,
        bold=True,
        color=RGBColor(0, 0, 0),
        after=4,
    )
    p.paragraph_format.space_before = Pt(20)
    add_para(
        doc,
        "面向不了解项目的读者：功能实现、技术选型、业务价值、实现难点与产品思考",
        size=13,
        color=RGBColor(55, 55, 55),
        after=14,
    )
    today = date.today().strftime("%Y-%m-%d")
    rows = [
        ("项目定位", "企业营销团队的 Multi-Agent AI 工作台"),
        ("核心用户", "市场营销、增长、内容、运营、管理者及需要营销洞察的业务团队"),
        ("当前版本", "具备登录、多用户隔离、会话管理、文件处理、实时执行追踪和产物交付能力的产品化原型"),
        ("技术栈", "Python / FastAPI / SQLite / Anthropic SDK / Next.js / React / TypeScript / Tailwind CSS"),
        ("生成日期", today),
    ]
    add_matrix(doc, ["字段", "说明"], rows, [1.35, 5.15])
    add_callout(
        doc,
        "一句话概括",
        "这个项目不是单纯的聊天机器人，而是把营销团队的内容生产、数据分析、市场研究和资讯监控流程，封装成一个可登录、可追踪、可沉淀结果的 AI 工作台。",
    )
    doc.add_page_break()


def add_overview(doc):
    add_heading(doc, "1. 项目总览", 1)
    add_para(
        doc,
        "Marketing Agent 面向企业营销团队，目标是把原本分散在文案工具、数据表格、搜索引擎、报告文档中的营销工作，整合到一个统一的 AI 工作台。用户只需要提出自然语言需求，系统会判断任务类型，分派给不同的专家 Agent，并把结果整理成可读、可追踪、可下载的交付物。",
    )
    add_heading(doc, "1.1 业务问题", 2)
    for item in [
        "营销团队的工作链路长：写内容、看数据、查竞品、做报告通常要在多个工具之间切换。",
        "AI 工具容易停留在单点生成：能写一段文案，但很难把研究、分析和内容产出串成完整业务流程。",
        "企业用户需要可追踪和可沉淀：AI 为什么这么输出、调用了哪些专家、产出了哪些文件，都需要能被回看。",
        "多用户场景需要隔离：不同账号的会话、上传文件、生成产物和新闻配置不能互相泄露。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "1.2 新版本相较早期版本的升级", 2)
    add_matrix(
        doc,
        ["维度", "早期形态", "新版本形态", "产品价值"],
        [
            ("交互入口", "偏 CLI 或单次 API 调用", "Next.js Web 工作台 + FastAPI API", "从技术 Demo 变成可体验的产品原型"),
            ("用户体系", "匿名或单用户", "注册、登录、资料维护、账号切换", "支持真实用户使用和数据归属"),
            ("会话能力", "临时上下文", "SQLite 持久化会话、消息、分组", "营销任务可沉淀、可复盘、可继续"),
            ("Agent 过程", "只返回最终结果", "SSE 展示调度、专家执行、Token 和产物事件", "提升透明度，降低黑箱感"),
            ("文件能力", "主要处理 CSV", "支持 CSV、Excel、JSON、PDF、Word、图片", "覆盖更真实的营销素材输入"),
            ("交付物", "Markdown 文本", "Markdown + PDF 产物 + 预览/下载", "更接近工作成果交付"),
            ("主动服务", "被动问答", "行业新闻摘要配置与刷新", "从工具型问答延伸到持续洞察服务"),
        ],
        [1.0, 1.6, 1.9, 2.0],
    )
    add_heading(doc, "1.3 面向用户的核心场景", 2)
    for item in [
        "内容生产：生成 LinkedIn、小红书、邮件、广告、博客、PDF 一页纸等营销内容。",
        "数据分析：上传投放或营销活动数据，自动计算 CTR、CVR、CPC、CPA、ROAS 等指标并给出建议。",
        "市场研究：查询趋势、竞品动作和行业信号，并要求带来源引用。",
        "行业新闻：配置关注行业、语言、时间和摘要粒度，生成近 24 小时新闻摘要。",
        "团队知识沉淀：通过会话、分组、历史记录和文件产物，沉淀一次营销任务的上下文与结果。",
    ]:
        add_bullet(doc, item)


def add_architecture(doc):
    add_heading(doc, "2. 系统架构与技术栈", 1)
    add_para(
        doc,
        "系统采用前后端分离架构。前端负责工作台体验、会话、上传、预览和执行链路展示；后端负责鉴权、数据持久化、文件管理、SSE 流式事件和 Agent 编排；AI 能力通过 Anthropic SDK 调用 Claude 模型及其服务端工具。",
    )
    add_matrix(
        doc,
        ["层级", "主要技术", "承担职责", "选择原因与收益"],
        [
            ("前端工作台", "Next.js 14、React 18、TypeScript、Tailwind CSS、lucide-react", "登录页、聊天区、会话侧边栏、文件上传、预览面板、执行追踪、主题切换", "适合快速构建 SaaS 型交互界面；TypeScript 提升接口稳定性；Tailwind 便于统一视觉规范"),
            ("API 层", "FastAPI、Uvicorn、sse-starlette", "REST 接口、SSE 流式响应、上传下载、会话读写、新闻配置", "FastAPI 开发效率高，类型清晰；SSE 比 WebSocket 更适合单向 AI 生成流"),
            ("持久化", "SQLite、WAL、外键、ContextVar", "用户、Token、会话、消息、分组、上传文件、产物、新闻配置和摘要", "MVP 阶段部署简单、零额外服务；外键和 user_id 支持基础数据隔离"),
            ("Agent 后端", "Python、Anthropic SDK、Claude Messages API", "Orchestrator 调度专家 Agent；专家 Agent 完成内容、分析、研究任务", "Python 生态适合文件处理、数据分析和后端服务；Anthropic 工具调用适合多 Agent 编排"),
            ("服务端工具", "web_search、code_execution、Files API、本地 generate_pdf", "联网研究、代码执行分析、文件上传进沙箱、PDF 生成", "让 Agent 不只生成文本，还能检索、计算和产出文件"),
            ("测试", "pytest / unittest、FastAPI TestClient、mock", "覆盖鉴权、会话、隔离、上传、产物、新闻、内容技能和分析 Agent", "降低核心链路回归风险，便于面试展示工程完整性"),
        ],
        [1.15, 1.6, 2.0, 1.75],
    )
    add_heading(doc, "2.1 数据和调用流", 2)
    for item in [
        "用户在前端输入需求并可附带文件。",
        "前端用 Bearer Token 调用 FastAPI；SSE 场景把 Token 放到查询参数中以兼容 EventSource。",
        "后端校验用户身份，确认 session、upload、artifact 等资源属于当前 user_id。",
        "后端构造用户消息：文本文件抽取内容，图片转成模型可识别的 image block，数据文件只传路径给分析 Agent。",
        "Orchestrator 先判断任务类型，再调用内容、分析或研究专家 Agent。",
        "专家 Agent 返回结果后，Orchestrator 综合输出；过程中通过 SSE 发送 delegating、specialist_done、assistant_delta、artifact_created 等事件。",
        "最终回答、错误信息和生成的文件产物会写入 SQLite，前端可回看、预览和下载。",
    ]:
        add_number(doc, item)
    add_callout(
        doc,
        "产品视角",
        "这个架构的重点不是“把模型接上去”，而是把模型能力放进真实工作流：有身份、有上下文、有文件、有过程、有结果沉淀。这样才更接近企业用户愿意长期使用的 AI 产品。",
    )


def add_feature_template(doc, title, business, implementation, tech, difficulties, rationale, benefits, improvements=None):
    add_heading(doc, title, 2)
    add_para(doc, "业务/产品目标", bold=True, color=DARK_BLUE, after=2)
    for item in business:
        add_bullet(doc, item)
    add_para(doc, "实现方式", bold=True, color=DARK_BLUE, after=2)
    for item in implementation:
        add_bullet(doc, item)
    add_para(doc, "使用技术", bold=True, color=DARK_BLUE, after=2)
    add_para(doc, "、".join(tech), after=4)
    add_para(doc, "可能遇到的困难", bold=True, color=DARK_BLUE, after=2)
    for item in difficulties:
        add_bullet(doc, item)
    add_para(doc, "为什么这样实现", bold=True, color=DARK_BLUE, after=2)
    for item in rationale:
        add_bullet(doc, item)
    add_para(doc, "带来的好处", bold=True, color=DARK_BLUE, after=2)
    for item in benefits:
        add_bullet(doc, item)
    if improvements:
        add_para(doc, "后续可优化方向", bold=True, color=DARK_BLUE, after=2)
        for item in improvements:
            add_bullet(doc, item)


def add_features(doc):
    add_heading(doc, "3. 功能模块详解", 1)
    add_feature_template(
        doc,
        "3.1 登录、注册与用户资料",
        [
            "让产品从匿名 Demo 变成真实用户可使用的工作台。",
            "为会话、文件、产物和新闻摘要建立清晰的数据归属。",
            "支持账号、头像、联系方式、公司和职位等资料维护，增强产品完整度。",
        ],
        [
            "注册时校验账号格式、密码、真实姓名、身份证号、头像和联系方式。",
            "密码使用 PBKDF2-HMAC-SHA256 加盐哈希后存储，不保存明文密码。",
            "登录成功后生成随机 token，写入 auth_tokens 表并设置有效期。",
            "前端把 token 存在 localStorage，并在普通 API 请求中放入 Authorization: Bearer 头。",
            "用户可更新资料、修改密码、退出登录或注销账号。",
        ],
        ["FastAPI", "SQLite", "secrets.token_urlsafe", "PBKDF2-HMAC-SHA256", "React localStorage"],
        [
            "鉴权链路一旦出错，会影响所有业务接口。",
            "localStorage 存 token 容易受 XSS 影响，需要后续加强前端安全策略。",
            "SSE 的 EventSource 不能自定义 Authorization Header，因此需要额外处理 token 传递。",
        ],
        [
            "MVP 阶段使用随机 token 比 JWT 更容易实现服务端失效、退出登录和删除 token。",
            "SQLite 存储 token 和用户表，部署简单，适合个人项目或面试作品快速展示完整链路。",
            "先实现可用的账号体系，有助于后续扩展团队、组织、权限和付费账户。",
        ],
        [
            "用户数据可以被归属和隔离。",
            "产品体验更像真实 SaaS，而不是临时工具。",
            "后续可基于账号做用量统计、权限控制、团队协作和商业化方案。",
        ],
        ["将 token 改为哈希后入库；引入 HttpOnly Cookie；增加登录限流、设备管理和 MFA。"],
    )

    add_feature_template(
        doc,
        "3.2 多用户数据隔离",
        [
            "确保不同用户的会话、上传文件、生成产物和新闻配置互不可见。",
            "为企业场景中的隐私、安全和合规要求打基础。",
        ],
        [
            "核心表 users、groups、sessions、uploads、artifacts、news_configs、news_summaries 均绑定 user_id。",
            "所有业务接口先调用 require_user 获取当前用户，再用 user_id 查询资源。",
            "访问 session、upload、artifact 时使用 id + user_id 双条件查询，跨用户访问返回 404。",
            "生成 PDF 产物时使用 ContextVar 绑定当前 user_id，确保 Agent 工具创建的 artifact 归属正确。",
            "删除用户时依赖外键级联删除相关 Token、会话、文件记录、新闻配置等数据。",
        ],
        ["SQLite 外键", "user_id 数据建模", "ContextVar", "FastAPI Request 鉴权"],
        [
            "Agent 工具可能在较深调用栈中创建产物，需要把用户上下文传递进去。",
            "文件物理存储在同一目录下，必须确保所有下载和预览入口都走数据库权限校验。",
            "跨用户访问需要统一返回 404，避免暴露资源是否存在。",
        ],
        [
            "MVP 阶段用 user_id 做逻辑隔离成本低，能快速覆盖最关键的隐私风险。",
            "外键级联让删除账号更可靠，避免残留孤儿数据。",
            "ContextVar 比到处传 user_id 更适合跨工具调用和异步流式处理场景。",
        ],
        [
            "满足多账号演示和基础企业安全预期。",
            "为后续组织空间、团队共享、RBAC 权限模型提供数据基础。",
            "让用户更放心上传真实营销数据和业务资料。",
        ],
        ["将文件物理目录也按 user_id 分区；增加对象存储签名 URL；完善审计日志。"],
    )

    add_feature_template(
        doc,
        "3.3 会话、消息和分组管理",
        [
            "让 AI 任务可以长期沉淀，而不是一次性对话。",
            "支持用户按项目、客户、活动或主题管理营销任务。",
        ],
        [
            "sessions 表保存会话名称、所属用户、所属分组、创建和更新时间。",
            "messages 表按 session_id 保存用户和助手消息，助手工具调用内容会被投影成可读文本。",
            "groups 表支持新建、重命名、删除分组，删除分组会级联删除其中会话。",
            "前端侧边栏支持新建会话、选择会话、重命名、移动分组、删除和右键菜单。",
            "服务端维护一个内存 Conversation cache，流式任务期间保留模型上下文，同时把消息持久化到 SQLite。",
        ],
        ["SQLite", "FastAPI", "React state", "自定义 sessions store", "Conversation cache"],
        [
            "模型需要完整上下文，而产品又需要持久化历史，两者的数据结构不同。",
            "工具调用消息如果原样展示会干扰用户，需要转成适合 UI 的文本。",
            "删除分组和删除会话要避免前端仍选中已删除资源。",
        ],
        [
            "SQLite 持久化保证服务重启后仍可恢复会话。",
            "内存 cache 让正在运行的 Agent 流程能直接追加上下文，减少频繁反序列化。",
            "分组管理符合营销团队按活动或客户组织工作的习惯。",
        ],
        [
            "用户可继续历史任务，减少重复描述背景的成本。",
            "项目更接近真实工作台，而不是一次性聊天窗口。",
            "会话数据可用于后续做搜索、复盘、知识库和协作功能。",
        ],
        ["增加全文搜索、收藏、归档、分享权限、会话模板和批量管理。"],
    )

    add_feature_template(
        doc,
        "3.4 聊天执行与 SSE 实时流式追踪",
        [
            "让用户看到 AI 正在做什么，而不是长时间等待一个未知结果。",
            "提升复杂 Agent 任务的透明度、可解释性和容错体验。",
        ],
        [
            "前端发送消息后创建 pending assistant 气泡，并建立 SSE 连接。",
            "后端通过 orchestrator_event_stream 把同步 Agent 回调桥接成异步事件流。",
            "事件包括 started、delegating、specialist_done、orchestrator_response、assistant_delta、artifact_created、error、cancelled。",
            "前端右侧 PreviewPanel 展示执行链路、专家状态、Token 输入输出、原始事件详情和产物预览。",
            "如果 SSE 断开，前端会尝试读取服务端已持久化消息；若仍无结果，再调用 complete 接口做非流式兜底。",
        ],
        ["Server-Sent Events", "sse-starlette", "FastAPI async iterator", "React EventSource 封装", "状态机 UI"],
        [
            "模型工具调用本身是非流式的，需要在拿到最终文本后再以小块 replay 的方式模拟打字效果。",
            "SSE 网络中断时，既不能丢结果，也不能重复写入过多消息。",
            "前端需要把不同事件映射成用户能理解的状态，而不是暴露底层技术噪音。",
        ],
        [
            "SSE 比 WebSocket 简单，适合服务端单向推送 AI 事件。",
            "把 Agent 过程拆成事件，有利于未来做可观测性、成本统计和调试。",
            "非流式兜底接口提升弱网络下的完成率。",
        ],
        [
            "用户等待感降低，信任度提升。",
            "产品能展示“AI 正在调度内容/分析/研究专家”，强化 Multi-Agent 差异化。",
            "技术上便于定位失败点，产品上便于解释结果来源。",
        ],
        ["接入真正的模型流式输出；增加任务取消、重试、步骤级耗时和成本估算。"],
    )

    add_feature_template(
        doc,
        "3.5 Orchestrator 与专家 Agent 编排",
        [
            "把用户模糊需求转成可执行的专家任务。",
            "避免一个模型同时承担所有角色导致职责混乱和输出不稳定。",
            "让产品能力可以按业务模块持续扩展。",
        ],
        [
            "Orchestrator 使用系统提示定义自己是企业营销团队 chief of staff。",
            "它只能通过 delegate_to_content_agent、delegate_to_analytics_agent、delegate_to_research_agent 三个工具分派任务。",
            "内容、分析、研究三个专家 Agent 各自有独立系统提示、工具和输出格式。",
            "如果任务需要多个专家，Orchestrator 被要求优先并行派发独立任务。",
            "专家返回后，Orchestrator 把多个结果综合成最终 Markdown 交付。",
        ],
        ["Anthropic Messages API", "Claude model", "Tool use", "Python 函数式 Agent 封装"],
        [
            "Orchestrator 可能越权自己写文案或编造外部事实，因此系统提示中设置硬规则。",
            "专家 Agent 失败时如果无限重试会造成成本和延迟失控，因此失败后不重复调用同一专家。",
            "多专家结果合成要避免信息冲突、重复和格式混乱。",
        ],
        [
            "主控 Agent + 专家 Agent 的结构符合复杂业务任务的分工逻辑。",
            "每个专家有清晰职责，有利于单独优化提示词、工具和模型。",
            "产品上可以把复杂 AI 能力解释成用户熟悉的“专家协作”。",
        ],
        [
            "提高输出稳定性和可维护性。",
            "便于后续新增 SEO Agent、品牌审校 Agent、投放优化 Agent 等角色。",
            "让执行追踪面板有明确的信息架构，可展示每个专家的进度。",
        ],
        ["为不同 Agent 配置不同模型；增加任务规划 JSON；引入专家结果评分和自动重试策略。"],
    )

    add_feature_template(
        doc,
        "3.6 内容生成与平台化写作技能",
        [
            "解决通用大模型文案容易“一个味道”的问题。",
            "让不同营销渠道的内容符合平台语境和格式约束。",
        ],
        [
            "content_skills.py 定义 LinkedIn、Twitter/X、小红书、Blog、Email、Ad Copy、PDF 等内容技能。",
            "每个技能包含 aliases、rules、output_contract、avoid 等规则。",
            "内容 Agent 根据用户任务、format 和 platform 选择最合适的技能，并把技能规则注入提示词。",
            "用户要求 PDF、一页纸或 brief 时，内容 Agent 调用本地 generate_pdf 工具生成文件。",
        ],
        ["Python dataclass", "规则化 Prompt 注入", "Anthropic tool use", "ReportLab PDF"],
        [
            "中文平台如小红书和英文平台如 LinkedIn 的语气差异很大，不能只靠一个通用提示词。",
            "内容格式需要可控，例如广告标题字数、邮件 Subject/Preheader、小红书段落和标签。",
            "PDF 生成需要把模型输出结构化成标题和章节，不能直接把长文本塞进文件。",
        ],
        [
            "用显式平台技能替代隐式提示词，有利于可维护和可测试。",
            "技能选择逻辑简单但有效，适合 MVP 先覆盖高频渠道。",
            "本地 PDF 工具让系统具备文件交付能力，不依赖前端截图或手动复制。",
        ],
        [
            "内容更贴近真实渠道，提高业务可用性。",
            "产品能展示“营销专家”而非普通写作助手。",
            "PDF 产物可下载、可预览，更接近用户实际交付场景。",
        ],
        ["加入品牌语调库、禁用词、审批流、A/B 文案评分和多语言本地化规则。"],
    )

    add_feature_template(
        doc,
        "3.7 营销数据分析 Agent",
        [
            "帮助营销人员从投放或活动数据中快速得到 KPI、发现和建议。",
            "降低非数据岗位处理 CSV/Excel 的门槛。",
        ],
        [
            "用户上传 CSV、Excel 或 JSON 后，后端保存文件并记录 user_id。",
            "数据文件不会被直接内联进 Prompt，而是传递本地路径给 Analytics Agent。",
            "Analytics Agent 使用 Anthropic Files API 上传文件，并以 container_upload 形式挂载到 code_execution 沙箱。",
            "模型在沙箱中用 pandas/numpy/openpyxl 读取数据、检查字段、计算 CTR、CVR、CPC、CPA、ROAS 等指标。",
            "输出包含 Key Metrics 表格、Findings 和 Recommendations。",
        ],
        ["Anthropic Files API", "code_execution_20260120", "pandas", "numpy", "openpyxl", "FastAPI UploadFile"],
        [
            "大文件直接进入 Prompt 会超过上下文，也会增加数据泄露和成本风险。",
            "用户数据字段可能不标准，模型需要先检查列名和类型，再决定可计算指标。",
            "分析结果必须避免编造，没有字段就要明确说明无法计算。",
        ],
        [
            "使用代码执行沙箱可以真正计算，而不是让模型凭文本猜测。",
            "container_upload 让原始数据不进入 Prompt，降低 token 成本并提升可处理文件规模。",
            "统一 KPI 输出格式符合营销团队阅读习惯。",
        ],
        [
            "把数据分析从“写公式/做透视表”变成自然语言任务。",
            "更容易形成可复用的营销分析报告模板。",
            "对产品经理来说，这是 AI 从生成内容走向辅助决策的关键能力。",
        ],
        ["增加字段映射 UI、指标解释、图表生成、异常检测和分析模板市场。"],
    )

    add_feature_template(
        doc,
        "3.8 市场研究 Agent 与来源引用",
        [
            "支持用户快速了解行业趋势、竞品动态和市场机会。",
            "缓解普通模型可能使用过期知识或编造事实的问题。",
        ],
        [
            "Research Agent 使用服务端 web_search 工具进行联网检索。",
            "系统提示要求收集 3-5 个可靠来源，尽量使用近 6 个月信息。",
            "输出必须包含 Summary、Findings、Implications for Marketing 和 Sources。",
            "每个事实性结论要求附 URL 和可见发布时间；推断需要明确标注为推断。",
            "新闻摘要功能复用 Research Agent，但把任务限制在近 24 小时窗口内。",
        ],
        ["Anthropic web_search tool", "Claude Messages API", "Prompt 约束", "来源引用输出格式"],
        [
            "搜索结果可能为空、过旧、格式异常或受工具额度限制。",
            "时间窗口新闻很容易混入旧新闻，需要在 Prompt 中明确发布时间边界。",
            "来源之间可能冲突，需要让模型说明差异而不是强行给单一结论。",
        ],
        [
            "服务端搜索工具比让用户手动贴链接更符合产品体验。",
            "强制引用让研究结果更可验证，适合企业场景。",
            "将事实和推断拆开，减少 AI 研究报告的误导风险。",
        ],
        [
            "让营销团队更快获得市场输入。",
            "为内容策略、竞品定位和活动规划提供依据。",
            "产品可信度明显高于无来源的普通聊天结果。",
        ],
        ["增加来源可信度评分、竞品监控、收藏来源、自动去重和定期推送。"],
    )

    add_feature_template(
        doc,
        "3.9 行业新闻摘要",
        [
            "让产品从被动问答扩展为主动洞察服务。",
            "帮助营销团队每天快速了解所在行业的最新变化。",
        ],
        [
            "用户在 NewsPanel 配置行业/主题、摘要时间、时区、语言和详细程度。",
            "配置存入 news_configs 表，每个用户一份配置。",
            "手动刷新时，后端读取配置并构造近 24 小时新闻检索任务。",
            "Research Agent 生成摘要后写入 news_summaries 表。",
            "如果研究失败，系统不会覆盖上一份有效摘要，避免用户看到空结果替代有效历史。",
        ],
        ["FastAPI", "SQLite", "ZoneInfo", "Research Agent", "React Markdown"],
        [
            "时间窗口需要处理用户所在时区，否则“近 24 小时”会不准确。",
            "搜索失败或无结果时，要保护用户已有摘要，不能让体验倒退。",
            "中英文语言要求需要贯穿标题、正文和分析，而不只是部分翻译。",
        ],
        [
            "配置化新闻摘要符合营销团队每日例会和行业监控习惯。",
            "复用 Research Agent 降低实现成本，也保持来源引用能力。",
            "失败不覆盖历史结果是一种产品级容错设计，比单纯报错更友好。",
        ],
        [
            "提高产品留存想象空间，因为用户可以形成每日打开习惯。",
            "把 AI Agent 从一次性工具变成持续服务。",
            "为未来自动化任务、提醒和订阅功能打基础。",
        ],
        ["接入后台定时调度、邮件/站内通知、摘要收藏、行业模板和多主题订阅。"],
    )

    add_feature_template(
        doc,
        "3.10 文件上传、抽取、预览与下载",
        [
            "让用户可以把真实业务素材带入 AI 工作流。",
            "让结果不仅停留在聊天文本，还能与文件、报告和素材联动。",
        ],
        [
            "上传接口限制文件大小，校验扩展名，保存到 tmp/uploads，并写入 uploads 表。",
            "支持 CSV、Excel、JSON、PDF、Word、PNG、JPG/JPEG。",
            "PDF 使用 pypdf 抽取文本；Word 使用 python-docx 抽取段落；图片转 base64 image block。",
            "数据文件不内联内容，只把路径交给 Analytics Agent 的代码执行沙箱。",
            "前端 PreviewPanel 支持 PDF iframe、图片预览、CSV 前 50 行表格预览，以及非可预览文件下载。",
        ],
        ["FastAPI UploadFile", "pypdf", "python-docx", "base64", "FileResponse", "React iframe/table/img"],
        [
            "不同文件类型处理方式差异大，尤其数据文件和文档文件不能用同一种 Prompt 策略。",
            "文件预览和下载必须走鉴权，否则可能造成越权访问。",
            "PDF/Word 抽取可能失败或丢失排版，需要在失败时给模型可理解的提示。",
        ],
        [
            "把上传、抽取、预览和下载分层实现，便于扩展更多文件类型。",
            "对数据文件采用路径/沙箱方式，对文档采用文本抽取，对图片采用视觉输入，分别匹配模型最擅长的处理方式。",
            "预览面板让用户无需离开工作台即可检查输入和产物。",
        ],
        [
            "覆盖真实营销工作中常见的素材输入。",
            "减少复制粘贴和工具切换，提高工作流完整度。",
            "支持更复杂的多模态任务，例如根据图片素材写文案、根据 PDF brief 生成方案。",
        ],
        ["增加病毒扫描、对象存储、OCR、PDF 表格抽取、图片压缩和上传进度条。"],
    )

    add_feature_template(
        doc,
        "3.11 产物管理与 PDF 交付",
        [
            "让 AI 的输出能成为可分享、可下载、可归档的工作成果。",
            "满足营销团队常见的一页纸、brief、方案 PDF 等交付需求。",
        ],
        [
            "内容 Agent 在需要文件交付时调用 generate_pdf 工具。",
            "generate_pdf 使用 ReportLab 渲染标题和章节，支持中文字体 fallback。",
            "生成后写入 artifacts 表，绑定 user_id；路由层收到 artifact_created 事件后再绑定 session_id。",
            "前端在聊天气泡中展示附件，并自动选中右侧预览面板。",
            "用户可通过 artifact preview/download 接口在线查看或下载 PDF。",
        ],
        ["ReportLab", "本地 client tool", "SQLite artifacts 表", "SSE artifact_created", "FastAPI FileResponse"],
        [
            "模型生成内容和文件生成是两个不同阶段，需要把模型的结构化参数转成稳定 PDF。",
            "Agent 工具创建文件时处于业务路由之外，用户归属和会话归属需要额外绑定。",
            "中文 PDF 需要处理字体，否则可能出现方块字或乱码。",
        ],
        [
            "本地 PDF 工具简单可控，适合 MVP 快速形成端到端产物。",
            "artifact 表把生成文件从聊天文本中独立出来，便于预览、下载和权限控制。",
            "通过 SSE 通知前端产物已生成，用户体验比刷新列表更自然。",
        ],
        [
            "AI 结果更接近真实业务交付物。",
            "用户可以把结果带出系统，用于汇报、分享或存档。",
            "为后续扩展 Word、PPT、图片海报等产物类型打基础。",
        ],
        ["增加 DOCX/PPTX 产物、模板系统、品牌样式、版本历史和在线编辑。"],
    )

    add_feature_template(
        doc,
        "3.12 前端工作台体验",
        [
            "把复杂的多 Agent 后端能力包装成用户能理解、能操作、能信任的界面。",
            "降低非技术用户使用 AI Agent 的门槛。",
        ],
        [
            "页面采用三栏布局：左侧会话/分组，中间聊天，右侧预览/追踪。",
            "支持侧边栏和预览栏折叠、拖拽宽度、深浅色主题、账号切换和用户菜单。",
            "ChatPanel 管理输入、附件、发送状态和消息气泡。",
            "PreviewPanel 在 preview 与 trace 两个 tab 之间切换，既能看文件，也能看执行过程。",
            "NewsPanel 独立承载行业新闻配置与摘要阅读，避免干扰主聊天场景。",
        ],
        ["Next.js App Router", "React Hooks", "TypeScript", "Tailwind CSS", "lucide-react", "react-markdown"],
        [
            "需要在一个页面同时管理鉴权、会话、文件、流式事件、预览和错误恢复，状态复杂。",
            "移动端和桌面端布局需要兼顾，但工作台类产品主要面向桌面高效操作。",
            "Agent 事件如果展示过多会让用户困惑，展示过少又失去透明度。",
        ],
        [
            "三栏工作台符合专业工具的扫描和多任务习惯。",
            "执行追踪和文件预览放在右侧，避免打断主聊天上下文。",
            "React + TypeScript 能让复杂状态和接口类型更可控。",
        ],
        [
            "用户能在一个界面完成需求输入、过程观察、结果阅读和文件下载。",
            "产品观感从“聊天框”升级为“营销 AI 操作台”。",
            "便于后续加入更多面板，如成本、任务队列、知识库和审批。",
        ],
        ["增加响应式移动体验、快捷指令、Prompt 模板、键盘快捷键和协作评论。"],
    )


def add_api_data(doc):
    add_heading(doc, "4. API 与数据模型概览", 1)
    add_heading(doc, "4.1 主要 API 模块", 2)
    add_matrix(
        doc,
        ["模块", "代表接口", "作用"],
        [
            ("鉴权", "/api/auth/register, /api/auth/login, /api/auth/me, /api/auth/logout", "用户注册、登录、资料读取/更新、退出和注销"),
            ("会话", "/api/sessions, /api/sessions/{id}/messages, /api/sessions/{id}/stream, /api/sessions/{id}/complete", "创建/管理会话、读取历史、流式执行和非流式兜底"),
            ("分组", "/api/groups", "会话分组的新建、重命名、删除和查询"),
            ("上传", "/api/upload, /api/uploads/{id}/preview, /api/uploads/{id}/download", "文件上传、预览和下载"),
            ("产物", "/api/artifacts/{id}, /api/artifacts/{id}/preview, /api/artifacts/{id}/download", "生成文件的元数据、预览和下载"),
            ("新闻", "/api/news/config, /api/news/summary, /api/news/refresh", "新闻摘要配置、读取最新摘要和手动刷新"),
        ],
        [1.1, 3.1, 2.3],
    )
    add_heading(doc, "4.2 数据表设计", 2)
    add_matrix(
        doc,
        ["表", "核心字段", "产品含义"],
        [
            ("users", "id, account, password_hash, username, profile fields", "账号主体，所有业务数据的归属源头"),
            ("auth_tokens", "token, user_id, expires_at", "登录态，支持服务端失效和退出登录"),
            ("groups", "id, user_id, name", "用户自己的会话分组"),
            ("sessions", "id, user_id, group_id, name, updated_at", "一次营销任务或持续对话"),
            ("messages", "session_id, role, content", "对话历史和 Agent 输出沉淀"),
            ("uploads", "id, user_id, original_name, mime, ext, path", "用户上传素材"),
            ("artifacts", "id, user_id, session_id, filename, mime, path", "AI 生成的文件产物"),
            ("news_configs", "user_id, industry, detail_level, summary_time, timezone, language", "用户的行业新闻订阅配置"),
            ("news_summaries", "user_id, summary, generated_at, window_start, window_end", "生成后的新闻摘要历史"),
        ],
        [1.25, 2.5, 2.75],
    )
    add_callout(
        doc,
        "数据建模的产品意义",
        "这些表不是单纯的后端实现细节。它们对应了 AI 产品从“即时生成”走向“用户资产沉淀”的关键能力：账号、任务、素材、过程、产物和长期配置。",
    )


def add_difficulties(doc):
    add_heading(doc, "5. 关键难点与解决思路", 1)
    add_matrix(
        doc,
        ["难点", "具体表现", "当前解决方式", "产品/业务收益"],
        [
            ("AI 黑箱感", "用户不知道系统是否在研究、分析还是写作", "SSE 执行链路面板展示调度和专家执行事件", "增强信任感和可解释性"),
            ("大文件分析", "CSV/Excel 直接进 Prompt 成本高且不稳定", "Files API + code_execution 沙箱读取文件", "支持更真实的数据分析场景"),
            ("多用户安全", "会话、文件、产物可能互相泄露", "所有资源绑定 user_id，接口按 user_id 查询", "满足基本隐私和企业使用预期"),
            ("专家职责混乱", "模型可能自己写文案、自己编数据或编事实", "Orchestrator 强制通过专家工具分派任务", "提升输出稳定性和业务可信度"),
            ("文件产物归属", "Agent 工具深层调用生成文件，难以绑定用户和会话", "ContextVar 绑定 user_id，路由事件再绑定 session_id", "生成文件可被安全预览、下载和回看"),
            ("搜索时效性", "行业新闻容易混入旧内容", "构造明确时间窗口和语言要求，失败不覆盖历史摘要", "新闻功能更可靠，避免用户看到错误更新"),
            ("前端状态复杂", "消息、SSE、附件、预览、会话同时变化", "分组件拆分 ChatPanel、PreviewPanel、SessionSidebar、NewsPanel", "降低交互复杂度，提升可维护性"),
        ],
        [1.1, 1.7, 1.9, 1.8],
    )
    add_heading(doc, "5.1 为什么这些选择适合当前阶段", 2)
    for item in [
        "SQLite 而不是 PostgreSQL：MVP 和面试作品重在展示完整产品闭环，SQLite 降低部署和维护成本；未来用户量提升后可平滑迁移。",
        "SSE 而不是 WebSocket：当前主要是服务端向前端推送 AI 执行事件，SSE 更简单、更轻量，浏览器原生支持。",
        "多 Agent 而不是单 Agent：营销任务天然分工明显，内容、数据、研究的工具、评估标准和提示词都不同，拆分后更容易优化。",
        "代码执行沙箱而不是让模型读表格文本：数据分析需要真实计算，沙箱方式更适合处理大文件和复杂指标。",
        "本地 PDF 工具而不是只返回 Markdown：营销工作强调交付物，文件产物能显著提升产品完成度。",
        "前端三栏工作台而不是普通聊天页：营销用户需要同时看历史、对话、文件和执行过程，工作台比单聊天框更符合高频业务操作。",
    ]:
        add_bullet(doc, item)


def add_business_pm(doc):
    add_heading(doc, "6. 产品经理视角总结", 1)
    add_heading(doc, "6.1 这个项目体现的 AI 产品能力", 2)
    for item in [
        "场景拆解能力：把企业营销工作拆成内容生产、数据分析、市场研究、资讯摘要和文件交付等模块。",
        "AI 工作流设计能力：不是停留在 Prompt，而是设计从输入、调度、执行、追踪到交付的完整链路。",
        "可信 AI 体验设计：通过来源引用、执行追踪、错误持久化、失败兜底和文件预览降低不确定性。",
        "数据与权限意识：从早期 Demo 直接升级到登录、多用户隔离、资源归属和删除级联。",
        "MVP 取舍能力：在可控复杂度下优先实现最能体现产品闭环的能力，而不是一开始追求大而全。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "6.2 适合在简历中表达的成果", 2)
    add_matrix(
        doc,
        ["简历表达", "背后的项目证据"],
        [
            ("设计并落地企业营销 Multi-Agent AI 工作台，覆盖内容生成、营销数据分析、市场研究和行业新闻摘要。", "Orchestrator + 内容/分析/研究专家 Agent；NewsPanel 与新闻配置/摘要表"),
            ("将 AI Agent Demo 产品化为 SaaS 原型，补齐登录、多用户隔离、会话管理、文件上传、执行追踪和产物预览下载。", "FastAPI 鉴权、SQLite user_id 隔离、Next.js 三栏工作台、SSE trace、artifacts"),
            ("围绕 AI 可信度设计执行透明化机制，实时展示 Agent 调度状态、专家结果、Token 消耗和产物生成过程。", "PreviewPanel trace tab、orchestrator_response usage、delegating/specialist_done/artifact_created 事件"),
            ("针对营销场景沉淀专家能力边界和输出规范，提升内容、数据、研究结果的业务可用性。", "content_skills、Analytics KPI Prompt、Research 来源引用和事实/推断区分"),
            ("引入多文件输入和代码执行沙箱，支持营销数据和业务素材进入 AI 工作流。", "uploads、file_inputs、Files API container_upload、code_execution"),
        ],
        [2.7, 3.8],
    )
    add_heading(doc, "6.3 后续产品路线图", 2)
    add_matrix(
        doc,
        ["阶段", "重点能力", "业务意义"],
        [
            ("短期", "不同 Agent 使用不同模型；任务取消/重试；更完整的错误提示和成本展示", "降低成本，提升稳定性和用户可控感"),
            ("中期", "品牌语调库、Prompt 模板、报告模板、图表生成、会话搜索", "让产品从通用 Agent 变成营销团队日常工具"),
            ("长期", "团队空间、权限体系、审批流、定时任务、CRM/广告平台数据接入", "进入企业协作和商业化场景"),
            ("商业化", "按席位、用量、产物模板、高级研究和数据连接器收费", "形成清晰的 SaaS 收费路径"),
        ],
        [1.0, 2.9, 2.6],
    )


def add_appendix(doc):
    add_heading(doc, "7. 附录：当前项目文件结构对应关系", 1)
    add_matrix(
        doc,
        ["目录/文件", "作用"],
        [
            ("src/marketing_agent/orchestrator.py", "主控 Agent，负责理解需求、分派专家和综合输出"),
            ("src/marketing_agent/agents/content_agent.py", "内容生成专家，支持平台技能和 PDF 产物"),
            ("src/marketing_agent/agents/analytics_agent.py", "数据分析专家，使用 Files API 和 code_execution 沙箱"),
            ("src/marketing_agent/agents/research_agent.py", "市场研究专家，使用 web_search 并要求来源引用"),
            ("src/marketing_agent/agents/content_skills.py", "平台化内容规则库"),
            ("server/routes.py", "FastAPI 核心路由：鉴权、会话、上传、产物、新闻、SSE"),
            ("server/db.py", "SQLite 数据模型和数据访问层"),
            ("server/auth.py", "账号校验、密码哈希、token 签发和鉴权"),
            ("server/news.py", "行业新闻摘要任务构造、生成和持久化"),
            ("server/uploads.py / file_inputs.py", "上传校验、文件保存、文本/图片抽取"),
            ("web/app/page.tsx", "前端工作台主页面和核心状态编排"),
            ("web/components/preview-panel.tsx", "右侧预览和执行链路面板"),
            ("web/components/session-sidebar.tsx", "会话和分组侧边栏"),
            ("web/components/news-panel.tsx", "行业新闻配置和摘要页面"),
            ("tests/", "核心接口、隔离、上传、产物、新闻、内容技能和分析 Agent 测试"),
        ],
        [2.65, 3.85],
    )
    add_callout(
        doc,
        "阅读建议",
        "不了解代码的读者可以先看第 1-3 章理解产品和功能，再看第 4-5 章理解技术实现与难点，最后用第 6 章提炼成面试或简历中的项目表达。",
    )


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    cover(doc)
    add_overview(doc)
    add_architecture(doc)
    add_features(doc)
    add_api_data(doc)
    add_difficulties(doc)
    add_business_pm(doc)
    add_appendix(doc)
    doc.save(OUT)
    print("DOCX_CREATED")


if __name__ == "__main__":
    build()
